# 导入必要的FastAPI和数据库相关模块
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form
from sqlalchemy.orm import Session
from app.database.connection import get_db
from app.services.embedding import embedding_service
from app.models.document import DocumentChunk
from app.models.database import Document, UserDocument, User
from app.models.schemas import DocumentResponse
from app.middleware.auth import get_current_active_user, require_document_upload, require_document_delete, require_document_read
from app.services.change_detector import ChangeDetector
from app.services.incremental_indexer import IncrementalIndexer
import PyPDF2
from docx import Document as DocxDocument
import json
import re
import hashlib
from datetime import datetime
import logging
from typing import List, Optional

# 创建路由器实例
router = APIRouter()
logger = logging.getLogger(__name__)


  # 推荐配置 1: 保守策略 (适合混合中英文)
#   MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
#   MAX_CHUNK_SIZE = 1000  # 约 500-750 tokens (中文), 750-1000 tokens (英文)
#   CHUNK_OVERLAP = 200    # 20% 重叠

#   # 推荐配置 2: 激进策略 (充分利用 ada-002 的能力)
#   MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB  
#   MAX_CHUNK_SIZE = 2000  # 约 1000-1500 tokens (中文), 1500-2000 tokens (英文)
#   CHUNK_OVERLAP = 400    # 20% 重叠

#   # 推荐配置 3: 技术文档优化
#   MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
#   MAX_CHUNK_SIZE = 1500  # 约 750-1125 tokens (中文)
#   CHUNK_OVERLAP = 300    # 20% 重叠



# 配置常量
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# 文档块切割参数
# 基于 text-embedding-ada-002 (8191 tokens) 优化
# 1500 字符 ≈ 750-1125 tokens (中文), 适合技术文档
MAX_CHUNK_SIZE = 1500  # 每个块的最大字符数
CHUNK_OVERLAP = 300    # 块之间的重叠字符数 (20% 重叠)

def clean_text_content(text: str) -> str:
    """
    清理文本内容,移除PostgreSQL不支持的字符

    Args:
        text: 原始文本

    Returns:
        清理后的文本
    """
    if not text:
        return text

    # 移除 NUL 字符 (\x00) - PostgreSQL 不允许
    text = text.replace('\x00', '')

    # 移除 Unicode 替换字符 (无效UTF-8编码的占位符)
    text = text.replace('\ufffd', '')

    # 移除其他控制字符(保留换行、制表符等常用字符)
    # 控制字符范围: \x00-\x1f (除了 \t \n \r)
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f]', '', text)

    return text

def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件中提取文本内容"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()

        # 清理提取的文本
        text = clean_text_content(text)
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise HTTPException(status_code=400, detail="Failed to extract text from PDF")

def extract_text_from_txt(file_path: str) -> str:
    """从文本文件中读取内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        # 清理读取的文本
        text = clean_text_content(text)
        return text
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        raise HTTPException(status_code=400, detail="Failed to read text file")

def extract_text_from_docx(file_path: str) -> str:
    """从DOCX文件中提取文本内容"""
    try:
        doc = DocxDocument(file_path)
        text = ""

        # 提取所有段落的文本
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        # 清理提取的文本
        text = clean_text_content(text)
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(status_code=400, detail="Failed to extract text from DOCX")

def calculate_content_hash(content: str) -> str:
    """计算文档内容的MD5哈希值"""
    return hashlib.md5(content.encode('utf-8')).hexdigest()

def split_text_into_chunks(text: str, chunk_size: int = MAX_CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    将文本分割成多个块

    Args:
        text (str): 要分割的文本
        chunk_size (int): 每个块的大小
        overlap (int): 块之间的重叠字符数

    Returns:
        List[str]: 分割后的文本块列表
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        # 确定当前块的结束位置
        end = start + chunk_size

        # 如果不是最后一块，尝试在句号、问号、感叹号处分割
        if end < len(text):
            # 查找最后一个句号、问号、感叹号
            last_sentence_end = text.rfind('。', start, end)
            if last_sentence_end == -1:
                last_sentence_end = text.rfind('？', start, end)
            if last_sentence_end == -1:
                last_sentence_end = text.rfind('！', start, end)
            if last_sentence_end == -1:
                last_sentence_end = text.rfind('.', start, end)
            if last_sentence_end == -1:
                last_sentence_end = text.rfind('?', start, end)
            if last_sentence_end == -1:
                last_sentence_end = text.rfind('!', start, end)

            # 如果找到句号，在句号后分割
            if last_sentence_end != -1 and last_sentence_end > start + chunk_size // 2:
                end = last_sentence_end + 1

        # 提取当前块
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # 计算下一个块的起始位置（考虑重叠）
        start = end - overlap
        if start >= len(text):
            break

    return chunks

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    namespace: str = Form('default'),  # 新增领域参数,从表单接收
    enable_change_detection: bool = Form(True),  # 是否启用变更检测
    db: Session = Depends(get_db),
    current_user: User = Depends(require_document_upload)
):
    """
    处理文档上传请求,支持PDF, TXT, DOCX格式,并可选启用变更检测

    【更新机制】:
    - 如果文档(filename + namespace)已存在,则检测内容变化
    - 内容未变化: 直接返回现有文档信息
    - 内容已变化: 删除旧chunks,更新文档,创建新chunks
    - 文档不存在: 创建新文档记录
    """
    try:
        # 验证文件类型
        if not file.filename.endswith(('.pdf', '.txt', '.docx', '.doc')):
            raise HTTPException(
                status_code=400,
                detail="Only PDF, TXT, DOC and DOCX files are supported"
            )

        # 读取文件内容
        file_content = await file.read()

        # 验证文件大小
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size allowed: {MAX_FILE_SIZE // (1024*1024)}MB"
            )

        file_path = f"/tmp/{file.filename}"

        # 保存临时文件
        with open(file_path, 'wb') as f:
            f.write(file_content)

        # 根据文件类型提取文本
        if file.filename.endswith('.pdf'):
            text_content = extract_text_from_pdf(file_path)
        elif file.filename.endswith(('.docx', '.doc')):
            text_content = extract_text_from_docx(file_path)
        else:
            text_content = extract_text_from_txt(file_path)

        # 验证文本内容(文本已在提取时清理过)
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No text content found in file")

        # 将文本分割成块
        text_chunks = split_text_into_chunks(text_content)
        logger.info(f"Document split into {len(text_chunks)} chunks")

        # 计算新内容的哈希值(用于变更检测)
        new_content_hash = calculate_content_hash(text_content)

        # ========== 【阶段1 核心修复】检查文档是否已存在 ==========
        existing_doc = db.query(Document).filter(
            Document.filename == file.filename,
            Document.namespace == namespace
        ).first()

        is_update = False
        main_document = None
        change_status = "new"

        if existing_doc:
            # ========== 文档已存在，检查内容是否变化 ==========
            logger.info(f"检测到已存在文档: id={existing_doc.id}, filename={file.filename}, namespace={namespace}")

            # 获取索引记录以对比哈希
            from app.models.index_record import DocumentIndexRecord
            existing_record = db.query(DocumentIndexRecord).filter(
                DocumentIndexRecord.doc_id == existing_doc.id
            ).first()

            if existing_record and existing_record.content_hash == new_content_hash:
                # ========== 内容未变化，直接返回 ==========
                logger.info(f"文档内容未变化 (hash={new_content_hash[:8]}...)")

                # 获取现有的chunk IDs
                existing_chunks = db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == existing_doc.id
                ).all()

                return {
                    "message": "Document unchanged - content is identical",
                    "document_id": existing_doc.id,
                    "document_chunk_ids": [chunk.id for chunk in existing_chunks],
                    "filename": file.filename,
                    "chunks_created": 0,
                    "total_chunks": len(existing_chunks),
                    "change_detection": {
                        "status": "unchanged",
                        "content_hash": new_content_hash,
                        "index_version": existing_record.index_version if existing_record else 0
                    }
                }

            # ========== 内容已变化，执行更新 ==========
            logger.info(f"文档内容已变化，开始更新: old_hash={existing_record.content_hash[:8] if existing_record else 'N/A'}..., new_hash={new_content_hash[:8]}...")

            is_update = True
            main_document = existing_doc
            change_status = "updated"

            # 【事务保护】在同一事务中删除旧chunks
            logger.info(f"删除旧的文档块...")
            deleted_chunks_count = db.query(DocumentChunk).filter(
                DocumentChunk.document_id == existing_doc.id
            ).delete(synchronize_session=False)
            logger.info(f"已删除 {deleted_chunks_count} 个旧文档块")

            # 更新主文档内容
            existing_doc.content = text_content
            existing_doc.doc_metadata = json.dumps({
                "filename": file.filename,
                "size": len(file_content),
                "type": file.filename.split('.')[-1],
                "total_chunks": len(text_chunks),
                "total_size": len(text_content),
                "user_id": current_user.id,
                "namespace": namespace,
                "updated_at": str(datetime.now())  # 记录更新时间
            })
            # created_at 保持不变，只更新内容

        else:
            # ========== 文档不存在，创建新记录 ==========
            logger.info(f"创建新文档: filename={file.filename}, namespace={namespace}")

            main_document = Document(
                content=text_content,
                embedding=None,
                doc_metadata=json.dumps({
                    "filename": file.filename,
                    "size": len(file_content),
                    "type": file.filename.split('.')[-1],
                    "total_chunks": len(text_chunks),
                    "total_size": len(text_content),
                    "user_id": current_user.id,
                    "namespace": namespace
                }),
                filename=file.filename,
                created_at=str(datetime.now()),
                namespace=namespace
            )

            db.add(main_document)
            db.commit()
            db.refresh(main_document)

            # 创建用户文档关联
            user_document = UserDocument(
                user_id=current_user.id,
                document_id=main_document.id,
                permission_level="write"
            )
            db.add(user_document)

        # ========== 创建文档块记录（新建或更新都需要）==========
        document_chunk_ids = []
        for i, chunk in enumerate(text_chunks):
            try:
                # 生成向量嵌入
                embedding = await embedding_service.create_embedding(chunk)

                # 创建文档块记录，关联到主文档
                document_chunk = DocumentChunk(
                    document_id=main_document.id,  # 复用existing_doc.id或新document.id
                    content=chunk,
                    embedding=embedding,
                    chunk_metadata=json.dumps({
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "chunk_size": len(chunk)
                    }),
                    chunk_index=i,
                    filename=f"{file.filename}_chunk_{i+1}",
                    created_at=str(datetime.now()),
                    namespace=namespace
                )

                db.add(document_chunk)
                db.flush()
                document_chunk_ids.append(document_chunk.id)
                db.commit()  # 提交事务

            except Exception as chunk_error:
                logger.error(f"Error processing chunk {i+1}: {chunk_error}")
                # 继续处理其他块
                continue

        if not document_chunk_ids:
            raise HTTPException(status_code=500, detail="Failed to process any chunks")

        # ========== 更新索引记录 ==========
        change_detection_result = None
        if enable_change_detection:
            try:
                from app.models.index_record import DocumentIndexRecord, IndexChangeHistory

                existing_record = db.query(DocumentIndexRecord).filter(
                    DocumentIndexRecord.doc_id == main_document.id
                ).first()

                now = datetime.now()

                if existing_record:
                    # 【阶段2：并发控制】使用乐观锁更新现有索引记录
                    old_hash = existing_record.content_hash
                    current_version = existing_record.index_version  # 记录当前版本号

                    # 使用 WHERE 子句检查版本号，实现乐观锁
                    update_count = db.query(DocumentIndexRecord).filter(
                        DocumentIndexRecord.doc_id == main_document.id,
                        DocumentIndexRecord.index_version == current_version  # 版本检查
                    ).update({
                        "content_hash": new_content_hash,
                        "chunk_count": len(document_chunk_ids),
                        "vector_count": len(document_chunk_ids),
                        "indexed_at": now,
                        "file_size": len(file_content),
                        "file_modified_at": now,
                        "index_version": current_version + 1  # 增加版本号
                    }, synchronize_session=False)

                    if update_count == 0:
                        # 版本号不匹配，说明有并发修改
                        db.rollback()
                        raise HTTPException(
                            status_code=409,
                            detail="Concurrent modification detected. The document was modified by another process. Please retry."
                        )

                    # 刷新记录以获取最新数据
                    db.refresh(existing_record)

                    change_detection_result = {
                        "status": change_status,  # "updated" or "new"
                        "old_hash": old_hash,
                        "new_hash": new_content_hash,
                        "index_version": existing_record.index_version
                    }

                    # 【阶段2】记录变更历史（仅更新时）
                    if change_status == "updated":
                        change_history = IndexChangeHistory(
                            doc_id=main_document.id,
                            change_type='content_modified',
                            old_hash=old_hash,
                            new_hash=new_content_hash,
                            changed_at=now,
                            change_metadata=json.dumps({
                                "updated_by": current_user.id,
                                "filename": file.filename,
                                "chunks_updated": len(document_chunk_ids),
                                "method": "implicit_upload"
                            })
                        )
                        db.add(change_history)
                else:
                    # 创建新索引记录
                    new_record = DocumentIndexRecord(
                        doc_id=main_document.id,
                        content_hash=new_content_hash,
                        chunk_count=len(document_chunk_ids),
                        vector_count=len(document_chunk_ids),
                        indexed_at=now,
                        file_size=len(file_content),
                        file_modified_at=now,
                        index_version=1,
                        namespace=namespace
                    )
                    db.add(new_record)

                    change_detection_result = {
                        "status": "new",
                        "content_hash": new_content_hash,
                        "index_version": 1
                    }

                    # 【阶段2】记录创建历史
                    change_history = IndexChangeHistory(
                        doc_id=main_document.id,
                        change_type='created',
                        old_hash=None,
                        new_hash=new_content_hash,
                        changed_at=now,
                        change_metadata=json.dumps({
                            "created_by": current_user.id,
                            "filename": file.filename,
                            "chunks_created": len(document_chunk_ids),
                            "method": "initial_upload"
                        })
                    )
                    db.add(change_history)

                db.commit()
                logger.info(f"变更检测完成: {change_detection_result}")

            except Exception as e:
                logger.error(f"变更检测失败: {e}")
                change_detection_result = {
                    "status": "error",
                    "error": str(e)
                }

        response = {
            "message": f"Document {'updated' if is_update else 'uploaded'} successfully",
            "document_id": main_document.id,
            "document_chunk_ids": document_chunk_ids,
            "filename": file.filename,
            "chunks_created": len(document_chunk_ids),
            "total_chunks": len(text_chunks),
            "is_update": is_update  # 新增字段，标识是否为更新操作
        }

        if change_detection_result:
            response["change_detection"] = change_detection_result

        return response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        db.rollback()  # 添加回滚保护
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

def _get_user_documents(db: Session, current_user: User, search_query: str = None):
    """
    获取用户有权限访问的文档列表(内部辅助函数)

    Args:
        db: 数据库会话
        current_user: 当前用户
        search_query: 可选的搜索关键词

    Returns:
        文档列表
    """
    from app.services.auth import auth_service
    is_admin = auth_service.has_permission(db, current_user, "user_management")

    if is_admin:
        # 管理员可以看到所有文档
        query = db.query(Document)
        if search_query:
            query = query.filter(
                Document.filename.contains(search_query) |
                Document.content.contains(search_query)
            )
        documents = query.all()
    else:
        # 普通用户只能看到自己的文档
        user_documents = db.query(UserDocument).filter(
            UserDocument.user_id == current_user.id
        ).all()
        document_ids = [ud.document_id for ud in user_documents]

        query = db.query(Document).filter(Document.id.in_(document_ids))
        if search_query:
            query = query.filter(
                Document.filename.contains(search_query) |
                Document.content.contains(search_query)
            )
        documents = query.all()

    return documents


def _format_document_response(doc: Document) -> DocumentResponse:
    """
    将Document模型转换为DocumentResponse(内部辅助函数)

    Args:
        doc: Document数据库模型

    Returns:
        DocumentResponse响应模型
    """
    return DocumentResponse(
        id=doc.id,
        filename=doc.filename,
        content=doc.content[:200] + "..." if len(doc.content) > 200 else doc.content,
        metadata=json.loads(doc.doc_metadata) if doc.doc_metadata else {},
        created_at=doc.created_at,
        namespace=doc.namespace if hasattr(doc, 'namespace') else 'default',
        domain_tags=doc.domain_tags if hasattr(doc, 'domain_tags') else {},
        domain_confidence=doc.domain_confidence if hasattr(doc, 'domain_confidence') else 0.0
    )


def _cascade_delete_document(db: Session, document_id: int) -> dict:
    """
    级联删除文档及其相关数据(内部辅助函数)

    Args:
        db: 数据库会话
        document_id: 文档ID

    Returns:
        删除统计信息字典
    """
    # 1. 删除所有文档块 (document_chunks)
    document_chunks = db.query(DocumentChunk).filter(
        DocumentChunk.document_id == document_id
    ).all()
    chunks_count = len(document_chunks)
    for chunk in document_chunks:
        db.delete(chunk)

    # 2. 删除所有用户文档关联 (user_documents)
    user_documents = db.query(UserDocument).filter(
        UserDocument.document_id == document_id
    ).all()
    associations_count = len(user_documents)
    for ud in user_documents:
        db.delete(ud)

    # 3. 删除主文档 (documents)
    document = db.query(Document).filter(Document.id == document_id).first()
    if document:
        db.delete(document)

    logger.info(f"文档 {document_id}: 删除了 {chunks_count} 个文档块, {associations_count} 个用户关联")

    return {
        "deleted_chunks": chunks_count,
        "deleted_user_associations": associations_count
    }


@router.get("/documents", response_model=list[DocumentResponse])
async def list_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(require_document_read)
):
    """获取所有已上传的文档列表"""
    try:
        documents = _get_user_documents(db, current_user)
        return [_format_document_response(doc) for doc in documents]
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to list documents")

@router.get("/documents/stats")
async def get_document_stats(db: Session = Depends(get_db)):
    """获取文档统计信息"""
    try:
        total_docs = db.query(Document).count()

        # 按文件类型统计
        docs_by_type = {}
        documents = db.query(Document).all()
        for doc in documents:
            metadata = json.loads(doc.doc_metadata) if doc.doc_metadata else {}
            file_type = metadata.get("file_type", "unknown")
            docs_by_type[file_type] = docs_by_type.get(file_type, 0) + 1

        # 最近7天的文档
        from datetime import datetime, timedelta
        recent_date = datetime.now() - timedelta(days=7)
        # 使用 cast 将字符串转换为日期时间进行比较
        from sqlalchemy import cast, DateTime
        recent_docs = db.query(Document).filter(
            cast(Document.created_at, DateTime) >= recent_date
        ).count()

        return {
            "total": total_docs,
            "by_type": docs_by_type,
            "recent": recent_docs
        }
    except Exception as e:
        logger.error(f"Error getting document stats: {e}")
        return {
            "total": 0,
            "by_type": {},
            "recent": 0
        }

@router.get("/documents/search", response_model=list[DocumentResponse])
async def search_documents(
    query: str = "",
    db: Session = Depends(get_db),
    current_user: User = Depends(require_document_read)
):
    """搜索文档(复用list_documents逻辑)"""
    try:
        if not query:
            return []

        # 复用 _get_user_documents 和 _format_document_response
        documents = _get_user_documents(db, current_user, search_query=query)
        return [_format_document_response(doc) for doc in documents]
    except Exception as e:
        logger.error(f"Error searching documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to search documents")

@router.get("/documents/{document_id}")
async def get_document(document_id: int, db: Session = Depends(get_db)):
    """获取单个文档详情"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        return DocumentResponse(
            id=document.id,
            filename=document.filename,
            content=document.content,
            metadata=json.loads(document.doc_metadata) if document.doc_metadata else {},
            created_at=document.created_at
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get document")

@router.put("/documents/{document_id}")
async def update_document(
    document_id: int,
    file: UploadFile = File(...),
    enable_change_detection: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_document_upload)
):
    """
    【阶段2】显式更新已存在的文档

    通过document_id直接更新文档内容
    - 检查文档是否存在
    - 检查用户权限
    - 对比内容哈希，如果未变化则不更新
    - 如果变化，删除旧chunks，创建新chunks
    - 更新索引记录
    - 记录变更历史
    """
    try:
        # 1. 检查文档是否存在
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # 2. 检查用户权限
        user_document = db.query(UserDocument).filter(
            UserDocument.user_id == current_user.id,
            UserDocument.document_id == document_id
        ).first()

        # 如果不是文档所有者，检查是否是管理员
        if not user_document:
            from app.services.auth import auth_service
            is_admin = auth_service.has_permission(db, current_user, "user_management")
            if not is_admin:
                raise HTTPException(status_code=403, detail="Permission denied - you don't own this document")

        # 3. 验证文件类型
        if not file.filename.endswith(('.pdf', '.txt', '.docx', '.doc')):
            raise HTTPException(400, "Only PDF, TXT, DOC and DOCX files are supported")

        # 4. 读取并验证文件内容
        file_content = await file.read()
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(413, f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB")

        # 5. 提取文本内容
        file_path = f"/tmp/{file.filename}"
        with open(file_path, 'wb') as f:
            f.write(file_content)

        if file.filename.endswith('.pdf'):
            text_content = extract_text_from_pdf(file_path)
        elif file.filename.endswith(('.docx', '.doc')):
            text_content = extract_text_from_docx(file_path)
        else:
            text_content = extract_text_from_txt(file_path)

        if not text_content.strip():
            raise HTTPException(400, "No text content found in file")

        # 6. 计算新内容哈希，检查是否真的变化
        new_content_hash = calculate_content_hash(text_content)

        from app.models.index_record import DocumentIndexRecord, IndexChangeHistory
        existing_record = db.query(DocumentIndexRecord).filter(
            DocumentIndexRecord.doc_id == document_id
        ).first()

        if existing_record and existing_record.content_hash == new_content_hash:
            # 内容未变化，直接返回
            logger.info(f"文档 {document_id} 内容未变化")
            existing_chunks = db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).all()

            return {
                "message": "Document content unchanged",
                "document_id": document_id,
                "filename": file.filename,
                "status": "unchanged",
                "chunks_count": len(existing_chunks),
                "change_detection": {
                    "status": "unchanged",
                    "content_hash": new_content_hash,
                    "index_version": existing_record.index_version
                }
            }

        # 7. 内容已变化，执行更新
        logger.info(f"文档 {document_id} 内容已变化，开始更新")

        text_chunks = split_text_into_chunks(text_content)
        logger.info(f"文档分割成 {len(text_chunks)} 个块")

        # 【事务保护】删除旧chunks
        deleted_count = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).delete(synchronize_session=False)
        logger.info(f"已删除 {deleted_count} 个旧文档块")

        # 8. 更新主文档
        document.content = text_content
        document.filename = file.filename  # 允许更新文件名
        document.doc_metadata = json.dumps({
            "filename": file.filename,
            "size": len(file_content),
            "type": file.filename.split('.')[-1],
            "total_chunks": len(text_chunks),
            "total_size": len(text_content),
            "user_id": current_user.id,
            "namespace": document.namespace,
            "updated_at": str(datetime.now()),
            "update_count": json.loads(document.doc_metadata).get("update_count", 0) + 1 if document.doc_metadata else 1
        })

        # 9. 创建新chunks
        document_chunk_ids = []
        for i, chunk in enumerate(text_chunks):
            try:
                embedding = await embedding_service.create_embedding(chunk)

                document_chunk = DocumentChunk(
                    document_id=document_id,
                    content=chunk,
                    embedding=embedding,
                    chunk_metadata=json.dumps({
                        "chunk_index": i,
                        "total_chunks": len(text_chunks),
                        "chunk_size": len(chunk)
                    }),
                    chunk_index=i,
                    filename=f"{file.filename}_chunk_{i+1}",
                    created_at=str(datetime.now()),
                    namespace=document.namespace
                )

                db.add(document_chunk)
                db.flush()
                document_chunk_ids.append(document_chunk.id)
                db.commit()

            except Exception as chunk_error:
                logger.error(f"处理块 {i+1} 失败: {chunk_error}")
                continue

        if not document_chunk_ids:
            raise HTTPException(500, "Failed to process any chunks")

        # 10. 更新索引记录（带乐观锁并发控制）
        change_detection_result = None
        if enable_change_detection:
            now = datetime.now()

            if existing_record:
                # 【阶段2：并发控制】使用乐观锁更新索引记录
                old_hash = existing_record.content_hash
                current_version = existing_record.index_version  # 记录当前版本号

                # 使用 WHERE 子句检查版本号，实现乐观锁
                update_count = db.query(DocumentIndexRecord).filter(
                    DocumentIndexRecord.doc_id == document_id,
                    DocumentIndexRecord.index_version == current_version  # 版本检查
                ).update({
                    "content_hash": new_content_hash,
                    "chunk_count": len(document_chunk_ids),
                    "vector_count": len(document_chunk_ids),
                    "indexed_at": now,
                    "file_size": len(file_content),
                    "file_modified_at": now,
                    "index_version": current_version + 1  # 增加版本号
                }, synchronize_session=False)

                if update_count == 0:
                    # 版本号不匹配，说明有并发修改
                    db.rollback()
                    raise HTTPException(
                        status_code=409,
                        detail="Concurrent modification detected. The document was modified by another process. Please retry."
                    )

                # 刷新记录以获取最新数据
                db.refresh(existing_record)

                change_detection_result = {
                    "status": "updated",
                    "old_hash": old_hash,
                    "new_hash": new_content_hash,
                    "index_version": existing_record.index_version  # 新版本号
                }
            else:
                # 如果没有索引记录，创建一个
                new_record = DocumentIndexRecord(
                    doc_id=document_id,
                    content_hash=new_content_hash,
                    chunk_count=len(document_chunk_ids),
                    vector_count=len(document_chunk_ids),
                    indexed_at=now,
                    file_size=len(file_content),
                    file_modified_at=now,
                    index_version=1,
                    namespace=document.namespace
                )
                db.add(new_record)

                change_detection_result = {
                    "status": "new_index",
                    "content_hash": new_content_hash,
                    "index_version": 1
                }

            # 11. 【阶段2】记录变更历史
            change_history = IndexChangeHistory(
                doc_id=document_id,
                change_type='content_modified',
                old_hash=existing_record.content_hash if existing_record else None,
                new_hash=new_content_hash,
                changed_at=now,
                change_metadata=json.dumps({
                    "updated_by": current_user.id,
                    "filename": file.filename,
                    "chunks_updated": len(document_chunk_ids),
                    "method": "explicit_put"
                })
            )
            db.add(change_history)

            db.commit()
            logger.info(f"文档 {document_id} 更新完成，变更已记录")

        return {
            "message": "Document updated successfully",
            "document_id": document_id,
            "filename": file.filename,
            "status": "updated",
            "chunks_created": len(document_chunk_ids),
            "chunks_deleted": deleted_count,
            "change_detection": change_detection_result
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"更新文档 {document_id} 失败: {e}")
        db.rollback()
        raise HTTPException(500, f"Update failed: {str(e)}")

@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_document_delete)
):
    """删除文档"""
    try:
        # 检查用户是否有权限删除此文档
        user_document = db.query(UserDocument).filter(
            UserDocument.user_id == current_user.id,
            UserDocument.document_id == document_id
        ).first()

        # 如果不是文档所有者，检查是否是管理员
        if not user_document:
            from app.services.auth import auth_service
            is_admin = auth_service.has_permission(db, current_user, "user_management")
            if not is_admin:
                raise HTTPException(status_code=403, detail="Permission denied")

        # 检查文档是否存在
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # 执行级联删除
        delete_stats = _cascade_delete_document(db, document_id)
        db.commit()

        return {
            "message": "Document deleted successfully",
            **delete_stats
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete document")

@router.post("/documents/batch")
async def delete_multiple_documents(
    document_ids: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_document_delete)
):
    """批量删除文档"""
    try:
        ids = document_ids.get("ids", [])
        if not ids:
            raise HTTPException(status_code=400, detail="No document IDs provided")

        # 检查是否是管理员
        from app.services.auth import auth_service
        is_admin = auth_service.has_permission(db, current_user, "user_management")

        # 统计变量
        total_deleted = 0
        total_chunks = 0
        total_associations = 0
        failed_ids = []

        for doc_id in ids:
            try:
                # 权限检查(非管理员只能删除自己的文档)
                if not is_admin:
                    user_document = db.query(UserDocument).filter(
                        UserDocument.user_id == current_user.id,
                        UserDocument.document_id == doc_id
                    ).first()
                    if not user_document:
                        logger.warning(f"用户 {current_user.id} 无权删除文档 {doc_id}")
                        failed_ids.append(doc_id)
                        continue

                # 检查文档是否存在
                document = db.query(Document).filter(Document.id == doc_id).first()
                if not document:
                    logger.warning(f"文档 {doc_id} 不存在")
                    failed_ids.append(doc_id)
                    continue

                # 执行级联删除
                delete_stats = _cascade_delete_document(db, doc_id)
                total_deleted += 1
                total_chunks += delete_stats["deleted_chunks"]
                total_associations += delete_stats["deleted_user_associations"]

            except Exception as e:
                logger.error(f"删除文档 {doc_id} 失败: {e}")
                failed_ids.append(doc_id)
                continue

        db.commit()

        result = {
            "message": f"Successfully deleted {total_deleted} documents",
            "total_deleted": total_deleted,
            "total_chunks_deleted": total_chunks,
            "total_associations_deleted": total_associations
        }

        if failed_ids:
            result["failed_ids"] = failed_ids
            result["failed_count"] = len(failed_ids)

        return result

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error batch deleting documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete documents")
